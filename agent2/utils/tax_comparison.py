import os
import json
import logging
import datetime
from pathlib import Path
from fpdf import FPDF
from PyPDF2 import PdfReader
import re
from openai import OpenAI
from dotenv import load_dotenv
from agent2.utils.tax_file_reader import read_tax_calculation_file

logger = logging.getLogger(__name__)
load_dotenv()

# Create output directory for reports
REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

def parse_previous_tax_return(file_obj):
    """
    Parse a previous year's tax return file (PDF, JSON, or DOCX).
    
    Args:
        file_obj: The uploaded file object
        
    Returns:
        dict: Extracted tax data or error message
    """
    try:
        # Check file extension
        filename = file_obj.name.lower()
        
        if filename.endswith(".json"):
            # Process JSON file
            file_obj.seek(0)  # Reset file pointer
            data = json.loads(file_obj.read().decode('utf-8'))
            return extract_tax_data_from_json(data)
        
        elif filename.endswith(".pdf"):
            # Process PDF file
            file_obj.seek(0)  # Reset file pointer
            return extract_tax_data_from_pdf(file_obj)
        
        elif filename.endswith(".docx"):
            # Process DOCX file
            file_obj.seek(0)  # Reset file pointer
            return extract_tax_data_from_docx(file_obj)
        
        else:
            return {"error": "Unsupported file format. Please upload a PDF, JSON, or DOCX file."}
    
    except Exception as e:
        logger.error(f"Error parsing previous tax return: {str(e)}")
        return {"error": f"Failed to parse tax return: {str(e)}"}

def extract_tax_data_from_json(json_data):
    """
    Extract relevant tax information from JSON data.
    
    Args:
        json_data (dict): The parsed JSON data
        
    Returns:
        dict: Extracted tax data
    """
    try:
        # For now, just return the data as is - we'll use AI to parse it later
        return {
            "source_type": "json",
            "raw_data": json_data
        }
    except Exception as e:
        logger.error(f"Error extracting data from JSON: {str(e)}")
        return {"error": f"Failed to extract data from JSON: {str(e)}"}

def extract_tax_data_from_pdf(file_obj):
    """
    Extract text content from a PDF file.
    
    Args:
        file_obj: The uploaded PDF file object
        
    Returns:
        dict: Extracted text content
    """
    try:
        # Save PDF to temporary file
        temp_path = REPORTS_DIR / f"temp_tax_return_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
        with open(temp_path, "wb") as f:
            f.write(file_obj.read())
            
        # Extract text from PDF
        text_content = ""
        pdf = PdfReader(temp_path)
        
        for page in pdf.pages:
            text_content += page.extract_text() + "\n"
        
        # Remove temp file
        os.remove(temp_path)
        
        return {
            "source_type": "pdf",
            "text_content": text_content,
            "page_count": len(pdf.pages)
        }
    except Exception as e:
        logger.error(f"Error extracting data from PDF: {str(e)}")
        return {"error": f"Failed to extract data from PDF: {str(e)}"}

def extract_tax_data_from_docx(file_obj):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_obj: The uploaded DOCX file object
        
    Returns:
        dict: Extracted text content
    """
    try:
        # Save DOCX to temporary file
        temp_path = REPORTS_DIR / f"temp_tax_return_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        with open(temp_path, "wb") as f:
            f.write(file_obj.read())
            
        # Extract text from DOCX
        from docx import Document
        doc = Document(temp_path)
        text_content = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " | "
                text_content += "\n"
        
        # Remove temp file
        os.remove(temp_path)
        
        return {
            "source_type": "docx",
            "text_content": text_content,
            "page_count": len(doc.paragraphs)  # Approximate measure
        }
    except Exception as e:
        logger.error(f"Error extracting data from DOCX: {str(e)}")
        return {"error": f"Failed to extract data from DOCX: {str(e)}"}

def generate_tax_comparison(previous_year_data, current_year_data=None, client_data=None):
    """
    Compare two tax documents and generate a PDF report.
    
    Args:
        previous_year_data (dict): Uploaded document data
        current_year_data (dict, optional): Baseline tax calculation. If None, will use base_tax_calculation.txt
        client_data (dict, optional): Client information
        
    Returns:
        dict: Result with PDF report path or error message
    """
    try:
        # If current_year_data not provided, use base_tax_calculation.txt
        if current_year_data is None:
            current_year_data = read_tax_calculation_file("base_tax_calculation.txt")
            if "error" in current_year_data:
                return current_year_data
            # Add source_type for consistency with previous_year_data
            current_year_data["source_type"] = "text"
            current_year_data["text_content"] = current_year_data.get("full_text", "")
        
        # If client_data is None, provide a minimal structure
        if client_data is None:
            client_data = {"name": "Tax Client", "tax_year": datetime.datetime.now().year}
        
        # Use OpenAI to analyze and compare the documents
        comparison_data = analyze_tax_returns_with_ai(
            previous_year_data, 
            current_year_data, 
            client_data
        )
        
        if "error" in comparison_data:
            return comparison_data
        
        # Generate PDF report
        pdf_path = create_comparison_report(comparison_data, client_data)
        
        return {"report_path": pdf_path, "comparison_data": comparison_data}
    
    except Exception as e:
        logger.error(f"Error generating document comparison: {str(e)}")
        return {"error": f"Failed to generate document comparison: {str(e)}"}

def analyze_tax_returns_with_ai(previous_year_data, current_year_data, client_data):
    """
    Use OpenAI to analyze and compare tax documents.
    
    Args:
        previous_year_data (dict): Uploaded document data
        current_year_data (dict): Baseline tax calculation data
        client_data (dict): Client information
        
    Returns:
        dict: AI analysis and structured comparison data
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logger.error("OpenAI API key not found")
        return {"error": "OpenAI API key not found. Please check your .env file."}
    
    try:
        # Format data based on source type for uploaded document
        if previous_year_data.get("source_type") == "json":
            # For JSON sources, handle potential nested dictionaries
            raw_data = previous_year_data.get("raw_data", {})
            
            # Check for nested dictionaries that might cause issues and flatten them
            flattened_data = {}
            for key, value in raw_data.items():
                if isinstance(value, dict):
                    # For dictionaries, convert to a string representation
                    flattened_data[key] = json.dumps(value)
                else:
                    flattened_data[key] = value
                    
            document1_str = json.dumps(flattened_data, indent=2)
        else:
            document1_str = previous_year_data.get("text_content", "")
        
        # Format baseline calculation data
        document2_str = current_year_data.get("full_text", "")
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create a simplified prompt for OpenAI to focus only on numeric value comparison
        prompt = f"""
        You are a tax expert analyzing two tax documents. Your task is to identify and compare ONLY numeric values 
        between these documents. Focus exclusively on extracting numeric tax metrics, amounts, and rates.

        # DOCUMENT 1 (Previous Year):
        ```
        {document1_str}
        ```

        # DOCUMENT 2 (Current Year):
        ```
        {document2_str}
        ```

        Please provide ONLY a detailed comparison of the numeric values found in both documents:
        
        1. DETAILED COMPARISON: Identify and compare ANY relevant tax metrics found in the documents. 
           - Extract numerical values from both documents where possible
           - Calculate differences
           - Do not include any analysis or recommendations
           - Strictly focus on numeric tax-related values only
        
        Generate structured JSON data for the comparison with this format:
        {{
          "year_labels": ["Previous Year", "Current Year"],
          "key_metrics": [
            {{
              "label": "[Tax Metric Name]",
              "document1": [numeric value],
              "document2": [numeric value],
              "difference": [numeric difference]
            }}
            // Include all relevant numeric metrics you can identify from the documents
          ]
        }}
        
        Include the JSON between markers [JSON_START] and [JSON_END] to make extraction easier.
        """
        
        # Call OpenAI API
        logger.info("Calling OpenAI for detailed tax document comparison...")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a tax expert who analyzes and compares tax documents, extracting and comparing only numeric metrics based on the actual content of the documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1  # Lower temperature for more consistent structured output
        )
        
        # Extract and parse response
        content = response.choices[0].message.content
        
        # Parse the response to get only the detailed comparison and JSON data
        analysis_sections = {
            "detailed_comparison": ""
        }
        
        # Extract detailed comparison section
        section_pattern = r"#*\s*(DETAILED COMPARISON):(.*?)(?=#|\[JSON_START\]|$)"
        section_matches = re.finditer(section_pattern, content, re.DOTALL)
        
        for match in section_matches:
            section_content = match.group(2).strip()
            analysis_sections["detailed_comparison"] = section_content
        
        # Extract JSON data
        json_start = content.find("[JSON_START]")
        json_end = content.find("[JSON_END]")
        
        comparison_data = {}
        
        if json_start >= 0 and json_end > json_start:
            json_content = content[json_start + 12:json_end].strip()
            try:
                comparison_data = json.loads(json_content)
                logger.info("Successfully extracted tax comparison metrics with AI")
            except json.JSONDecodeError:
                logger.error("Failed to parse JSON from OpenAI response")
                
                # Attempt to extract JSON using regex as fallback
                json_pattern = r'{\s*"year_labels":\s*\[.*?\],\s*"key_metrics":\s*\[.*?\]\s*}'
                match = re.search(json_pattern, content, re.DOTALL)
                if match:
                    try:
                        comparison_data = json.loads(match.group(0))
                        logger.info("Successfully extracted tax comparison metrics with fallback regex")
                    except:
                        pass
        else:
            # If JSON markers aren't found, try to extract using regex
            json_pattern = r'{\s*"year_labels":\s*\[.*?\],\s*"key_metrics":\s*\[.*?\]\s*}'
            match = re.search(json_pattern, content, re.DOTALL)
            if match:
                try:
                    comparison_data = json.loads(match.group(0))
                    logger.info("Successfully extracted tax comparison metrics with regex")
                except:
                    logger.error("Failed to parse JSON from regex match")
        
        # If we still don't have valid JSON data, create minimal structure
        if not comparison_data or "key_metrics" not in comparison_data:
            comparison_data = {
                "year_labels": ["Document 1", "Document 2"],
                "key_metrics": []
            }
            
        # Combine analysis sections and comparison data
        result = {
            "year_labels": comparison_data.get("year_labels", ["Document 1", "Document 2"]),
            "key_metrics": comparison_data.get("key_metrics", []),
            "analysis": analysis_sections,
            "full_analysis_text": content
        }
        
        return result
    
    except Exception as e:
        logger.error(f"Error analyzing tax returns with AI: {str(e)}")
        return {"error": f"Failed to analyze tax returns: {str(e)}"}

def create_comparison_report(comparison_data, client_data):
    """
    Create a comprehensive PDF report of the tax document comparison.
    
    Args:
        comparison_data (dict): The comparison data from AI analysis
        client_data (dict): Client information
        
    Returns:
        str: Path to the generated PDF report
    """
    try:
        # Generate a PDF report with FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Add title page
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Tax Document Comparison Report", 0, 1, "C")
        
        # Add client name if available
        client_name = client_data.get("name", "")
        if not client_name and "ClientDetails" in client_data:
            client_name = client_data["ClientDetails"].get("name", "")
        
        if client_name:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, f"For: {client_name}", 0, 1, "C")
        
        # Add date
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}", 0, 1, "C")
        
        # Add some spacing between title and comparison data
        pdf.ln(5)
        
        # Add comparison data on the same page - removed the pdf.add_page() call that was here before
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Tax Metrics Comparison", 0, 1, "C")
        
        # Create table headers
        pdf.set_fill_color(220, 220, 220)
        pdf.set_font("Arial", "B", 10)
        
        # Get document labels
        doc_labels = comparison_data.get("year_labels", ["Previous Year", "Current Year"])
        
        # Set column widths for better readability - adjusted for three columns
        metric_width = 70
        doc1_width = 45
        doc2_width = 45
        diff_width = 30
        
        # Draw table header for three columns
        pdf.cell(metric_width, 10, "Tax Metric", 1, 0, "L", True)
        pdf.cell(doc1_width, 10, doc_labels[0], 1, 0, "C", True)
        pdf.cell(doc2_width, 10, doc_labels[1], 1, 0, "C", True)
        pdf.cell(diff_width, 10, "Difference", 1, 1, "C", True)  # End the row (1,1)
        
        # Add data rows
        pdf.set_font("Arial", "", 10)
        
        # Keep track of alternating row colors for readability
        row_count = 0
        
        # Process each metric identified by OpenAI
        for metric in comparison_data.get("key_metrics", []):
            # Alternate row colors
            if row_count % 2 == 0:
                pdf.set_fill_color(245, 245, 245)
                has_fill = True
            else:
                has_fill = False
            row_count += 1
            
            # Extract values - handle different key names
            label = metric.get("label", "")
            
            # Handle potential different keys in the metrics
            doc1_value = metric.get("document1", metric.get("previous_year", metric.get("doc1", 0)))
            doc2_value = metric.get("document2", metric.get("current_year", metric.get("doc2", 0)))
            
            # Calculate difference if not provided
            if "difference" in metric:
                diff = metric["difference"]
            else:
                diff = doc2_value - doc1_value if isinstance(doc2_value, (int, float)) and isinstance(doc1_value, (int, float)) else 0
            
            # Render label cell
            pdf.cell(metric_width, 8, label, 1, 0, "L", has_fill)
            
            # Format and render document 1 value
            if isinstance(doc1_value, (int, float)) and "rate" not in label.lower():
                pdf.cell(doc1_width, 8, f"${doc1_value:,.2f}", 1, 0, "R", has_fill)
            elif isinstance(doc1_value, (int, float)) and "rate" in label.lower():
                pdf.cell(doc1_width, 8, f"{doc1_value:.2f}%", 1, 0, "R", has_fill)
            else:
                pdf.cell(doc1_width, 8, str(doc1_value), 1, 0, "R", has_fill)
            
            # Format and render document 2 value
            if isinstance(doc2_value, (int, float)) and "rate" not in label.lower():
                pdf.cell(doc2_width, 8, f"${doc2_value:,.2f}", 1, 0, "R", has_fill)
            elif isinstance(doc2_value, (int, float)) and "rate" in label.lower():
                pdf.cell(doc2_width, 8, f"{doc2_value:.2f}%", 1, 0, "R", has_fill)
            else:
                pdf.cell(doc2_width, 8, str(doc2_value), 1, 0, "R", has_fill)
            
            # Format difference - this is now the last column
            if isinstance(diff, (int, float)):
                sign = "+" if diff > 0 else ""
                if "rate" in label.lower():
                    diff_str = f"{sign}{diff:.2f}%"
                else:
                    diff_str = f"{sign}${diff:,.2f}"
                    
                pdf.cell(diff_width, 8, diff_str, 1, 1, "R", has_fill)  # End the row (1,1)
            else:
                pdf.cell(diff_width, 8, "N/A", 1, 1, "R", has_fill)  # End the row (1,1)
        
        # Save the PDF
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client_name = "".join(c if c.isalnum() else "_" for c in client_name) if client_name else "tax_client"
        pdf_path = str(REPORTS_DIR / f"{safe_client_name}_Tax_Comparison_{timestamp}.pdf")
        
        pdf.output(pdf_path)
        logger.info(f"Generated numeric tax comparison report: {pdf_path}")
        
        return pdf_path
        
    except Exception as e:
        logger.error(f"Error creating comparison report: {str(e)}")
        raise Exception(f"Failed to create comparison report: {str(e)}")
